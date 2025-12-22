import os
import re
import uuid
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple, Union
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from pathlib import Path
import json
from loguru import logger
import matplotlib
# 设置matplotlib后端为Agg，必须在导入pyplot之前设置
matplotlib.use('Agg')
logger.info(f"Matplotlib backend set to: {matplotlib.get_backend()}")

from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib import rcParams
import seaborn as sns
import pandas as pd
from io import BytesIO
import base64

from app.core.config import settings
from app.models.analysis import AuditReport
from app.models.project import Project, ProjectMaterial, ProjectType
from app.models.analysis import PriceAnalysis
from app.models.user import User


PROVINCE_MAP = {
    '110000': '北京市', '120000': '天津市', '130000': '河北省', '140000': '山西省',
    '150000': '内蒙古自治区', '210000': '辽宁省', '220000': '吉林省', '230000': '黑龙江省',
    '310000': '上海市', '320000': '江苏省', '330000': '浙江省', '340000': '安徽省',
    '350000': '福建省', '360000': '江西省', '370000': '山东省', '410000': '河南省',
    '420000': '湖北省', '430000': '湖南省', '440000': '广东省', '450000': '广西壮族自治区',
    '460000': '海南省', '500000': '重庆市', '510000': '四川省', '520000': '贵州省',
    '530000': '云南省', '540000': '西藏自治区', '610000': '陕西省', '620000': '甘肃省',
    '630000': '青海省', '640000': '宁夏回族自治区', '650000': '新疆维吾尔自治区'
}

CITY_MAP = {
    '110100': '北京市', '120100': '天津市', '130100': '石家庄市', '140100': '太原市',
    '150100': '呼和浩特市', '210100': '沈阳市', '220100': '长春市', '230100': '哈尔滨市',
    '310100': '上海市', '320100': '南京市', '330100': '杭州市', '340100': '合肥市',
    '350100': '福州市', '360100': '南昌市', '370100': '济南市', '410100': '郑州市',
    '420100': '武汉市', '430100': '长沙市', '440100': '广州市', '450100': '南宁市',
    '460100': '海口市', '500100': '重庆市', '510100': '成都市', '520100': '贵阳市',
    '530100': '昆明市', '540100': '拉萨市', '610100': '西安市', '620100': '兰州市',
    '630100': '西宁市', '640100': '银川市', '650100': '乌鲁木齐市', '320500': '苏州市',
    '330200': '宁波市', '440300': '深圳市'
}

DISTRICT_MAP = {
    '330102': '上城区', '330104': '拱墅区', '330105': '西湖区', '330106': '滨江区',
    '330107': '萧山区', '330108': '余杭区', '330109': '富阳区', '330110': '余杭区',
    '330111': '富阳区', '330112': '临安区', '330113': '临平区', '330114': '钱塘区',
    '330122': '桐庐县', '330127': '淳安县', '330182': '建德市', '330101': '上城区',
    '330103': '江干区'
}

AUDIT_SCOPE_MAP = {
    'price_analysis': '价格合理性分析',
    'material_matching': '材料匹配度检查',
    'market_comparison': '市场价格对比',
    'risk_assessment': '风险评估'
}


class ReportGenerator:
    """审计报告生成器"""
    
    def __init__(self):
        self.templates_dir = Path("templates")
        
        # 使用绝对路径，确保在不同环境下都能正常工作
        if hasattr(settings, 'REPORTS_DIR') and settings.REPORTS_DIR:
            self.reports_dir = Path(settings.REPORTS_DIR)
        else:
            # 使用当前工作目录下的reports文件夹
            self.reports_dir = Path.cwd() / "reports"
            
        self.reports_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"ReportGenerator output directory: {self.reports_dir.absolute()}")
        
        # 配置matplotlib中文支持（增加常见中文字体作为回退）
        rcParams['font.family'] = 'sans-serif'
        rcParams['font.sans-serif'] = [
            'Arial Unicode MS',
            'Heiti TC',
            'Songti SC',
            'STHeiti',
            'Microsoft YaHei',
            'SimSun',
            'PingFang SC',
            'Noto Sans CJK SC',
            'SimHei'
        ]
        rcParams['axes.unicode_minus'] = False
        
    async def generate_audit_report(
        self,
        db: AsyncSession,
        project: Project,
        materials: List[ProjectMaterial],
        analyses: List[PriceAnalysis],
        report_config: Dict[str, Any] = None,
        chart_images: Dict[str, str] = None
    ) -> str:
        """生成完整的审计报告
        
        Args:
            project: 项目信息
            materials: 项目材料列表
            analyses: 价格分析结果
            report_config: 报告配置参数
            chart_images: 前端传入的图表图片(base64)
            
        Returns:
            报告文件路径
        """
        try:
            # 解析配置
            config = report_config or {}
            include_charts = config.get('include_charts', True)
            include_details = config.get('include_detailed_analysis', True)
            include_recommendations = config.get('include_recommendations', True)
            include_appendices = config.get('include_appendices', True)

            # 生成报告数据
            report_data = await self._prepare_report_data(db, project, materials, analyses)

            # 创建Word文档
            doc = Document()
            self._setup_document_styles(doc)
            self._apply_document_layout(doc)

            self._add_header_footer(doc, report_title='造价材料审计报告')
            # 添加报告标题页
            self._add_title_page(doc, report_data)

            # 添加执行摘要
            self._add_executive_summary(doc, report_data)

            # 添加项目概况
            self._add_project_overview(doc, report_data)

            # 添加分析方法与数据来源
            self._add_methodology_section(doc, report_data)

            # 添加分析结果
            chart_files = []
            if include_charts:
                try:
                    if chart_images:
                        chart_files = self._process_chart_images(chart_images)
                    else:
                        chart_files = self._generate_charts(report_data)
                except Exception as e:
                    logger.error(f"生成图表失败: {e}")

            await self._add_analysis_results(doc, report_data, include_details=include_details, chart_files=chart_files)

            # 添加问题材料详情
            # self._add_problematic_materials(doc, report_data)

            # 添加建议措施
            if include_recommendations:
                self._add_recommendations(doc, report_data)

            # 添加附录
            if include_appendices:
                self._add_appendices(doc, report_data)

            # 保存报告
            # 使用项目名称作为文件名的一部分
            project_name_safe = re.sub(r'[\\/*?:"<>|]', '_', project.name) if project.name else f"project_{project.id}"
            report_filename = f"分析报告_{project_name_safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            report_path = self.reports_dir / report_filename
            doc.save(str(report_path))

            logger.info(f"审计报告生成成功: {report_path}")
            return str(report_path)
            
        except Exception as e:
            logger.error(f"生成审计报告失败: {e}")
            raise
    
    async def _prepare_report_data(
        self,
        db: AsyncSession,
        project: Project,
        materials: List[ProjectMaterial],
        analyses: List[PriceAnalysis]
    ) -> Dict[str, Any]:
        """准备报告数据"""

        total_materials = len(materials)

        price_analyses = {a.material_id: a for a in analyses}

        analyzed_ids_from_flag = {m.id for m in materials if getattr(m, 'is_analyzed', False)}
        analyzed_ids_from_analysis = {a.material_id for a in analyses if a is not None}
        analyzed_materials = len(analyzed_ids_from_flag | analyzed_ids_from_analysis)

        problematic_from_material = {m.id for m in materials if getattr(m, 'is_problematic', False)}
        problematic_from_analysis = {a.material_id for a in analyses if a and a.is_reasonable is False}
        problematic_materials = len(problematic_from_material | problematic_from_analysis)
        unreasonable_count = len(problematic_from_analysis)

        def infer_risk_level(a: Optional[PriceAnalysis]) -> str:
            if not a:
                return 'unknown'
            if a.risk_level:
                return a.risk_level
            try:
                variance = abs(a.price_variance or 0)
                if variance >= 60:
                    return 'critical'
                if variance >= 40:
                    return 'high'
                if variance >= 20:
                    return 'medium'
                if variance > 0:
                    return 'low'
                return 'normal'
            except Exception:
                return 'unknown'

        risk_stats: Dict[str, int] = {}
        for analysis in analyses:
            level = infer_risk_level(analysis)
            risk_stats[level] = risk_stats.get(level, 0) + 1

        price_variances = [a.price_variance for a in analyses if a and a.price_variance is not None]
        avg_variance = sum(price_variances) / len(price_variances) if price_variances else 0

        estimated_savings = 0
        for material in materials:
            analysis = price_analyses.get(material.id)
            if analysis and analysis.predicted_price_avg and material.unit_price:
                if material.unit_price > analysis.predicted_price_avg:
                    savings = (material.unit_price - analysis.predicted_price_avg) * (material.quantity or 0)
                    estimated_savings += savings

        # 获取创建人信息
        creator_name = "未知"
        if project.created_by:
            try:
                result = await db.execute(select(User).where(User.id == project.created_by))
                user = result.scalar_one_or_none()
                if user:
                    creator_name = user.full_name or user.username
            except Exception:
                pass

        from app.services.report_service import ReportService
        report_service = ReportService()

        analysis_materials_raw = report_service._generate_analysis_materials_data(materials, analyses)
        guidance_materials_raw = await report_service._generate_guidance_price_materials_data(
            db, project.id, materials, analyses
        )

        def transform_analysis_items(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
            transformed: List[Dict[str, Any]] = []
            for item in items:
                quantity = float(item.get('quantity') or 0)
                original_unit = float(item.get('original_price') or 0)
                ai_unit = float(item.get('ai_predicted_price') or item.get('predicted_price') or 0)
                original_total = float(item.get('original_total_price') or (original_unit * quantity))
                ai_total = float(item.get('ai_total_price') or item.get('ai_total') or (ai_unit * quantity))
                adjustment = float(item.get('adjustment')) if item.get('adjustment') is not None else original_total - ai_total
                risk_level = (item.get('risk_level') or '').lower()
                is_reasonable = item.get('is_reasonable')

                has_difference = abs(adjustment) > 1e-2
                has_risk = (risk_level not in ('', 'normal', 'low')) or (is_reasonable is False)
                if not (has_difference or has_risk):
                    continue

                transformed.append({
                    'materialName': item.get('material_name', ''),
                    'specification': item.get('specification', ''),
                    'unit': item.get('unit', ''),
                    'quantity': quantity,
                    'originalUnitPrice': original_unit,
                    'originalTotalPrice': original_total,
                    'aiUnitPrice': ai_unit,
                    'aiTotalPrice': ai_total,
                    'adjustment': adjustment,
                    'weightPercentage': 0.0,
                    'riskLevel': risk_level or 'normal'
                })

            total_original = sum(abs(item['originalTotalPrice']) for item in transformed)
            if total_original > 0:
                for item in transformed:
                    item['weightPercentage'] = (abs(item['originalTotalPrice']) / total_original) * 100

            return transformed

        def transform_guidance_items(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
            transformed: List[Dict[str, Any]] = []
            for item in items:
                if item.get('material_name') in ['暂无市场信息价材料数据', '数据获取失败', None, '']:
                    continue

                quantity = float(item.get('quantity') or 0)
                original_unit = float(item.get('original_price') or 0)
                guidance_unit = float(item.get('guidance_price') or item.get('base_price') or 0)
                original_total = float(item.get('original_total_price') or (original_unit * quantity))
                guidance_total = float(item.get('guidance_total_price') or item.get('ai_total_price') or (guidance_unit * quantity))
                adjustment = float(item.get('adjustment')) if item.get('adjustment') is not None else original_total - guidance_total

                if abs(adjustment) <= 1e-2:
                    continue

                transformed.append({
                    'materialName': item.get('material_name', ''),
                    'specification': item.get('specification', ''),
                    'unit': item.get('unit', ''),
                    'quantity': quantity,
                    'originalUnitPrice': original_unit,
                    'originalTotalPrice': original_total,
                    'aiUnitPrice': guidance_unit,
                    'aiTotalPrice': guidance_total,
                    'adjustment': adjustment,
                    # 新增字段，与前端 GuidancePriceMaterialTable 保持一致
                    'originalBasePrice': float(item.get('original_base_price') or 0),
                    'priceDiff': float(item.get('price_diff') or 0),
                    'riskRate': float(item.get('risk_rate') or 0),
                    'weightPercentage': 0.0
                })

            total_original = sum(abs(item['originalTotalPrice']) for item in transformed)
            if total_original > 0:
                for item in transformed:
                    item['weightPercentage'] = (abs(item['originalTotalPrice']) / total_original) * 100

            return transformed

        analysis_materials_report = transform_analysis_items(analysis_materials_raw)
        guidance_materials_report = transform_guidance_items(guidance_materials_raw)

        def calc_totals(items: List[Dict[str, Any]]) -> Dict[str, float]:
            return {
                'original_total': sum(item['originalTotalPrice'] for item in items),
                'ai_total': sum(item['aiTotalPrice'] for item in items),
                'adjustment_total': sum(item['adjustment'] for item in items)
            }

        analysis_totals = calc_totals(analysis_materials_report)
        guidance_totals = calc_totals(guidance_materials_report)

        combined_adjustments = []
        for item in analysis_materials_report:
            combined_adjustments.append((item['materialName'], item['adjustment']))
        for item in guidance_materials_report:
            combined_adjustments.append((item['materialName'], item['adjustment']))

        combined_adjustments.sort(key=lambda x: abs(x[1]), reverse=True)
        top_adjustments = combined_adjustments[:5]

        risk_label_map = {
            'normal': '正常',
            'low': '低风险',
            'medium': '中风险',
            'high': '高风险',
            'critical': '极高风险',
            'severe': '极高风险',
            'unknown': '待确认'
        }
        risk_distribution = {
            risk_label_map.get(level, level): count
            for level, count in risk_stats.items()
        }

        return {
            'project': project,
            'db': db,
            'report_date': datetime.now(),
            'statistics': {
                'total_materials': total_materials,
                'analyzed_materials': analyzed_materials,
                'problematic_materials': problematic_materials,
                'unreasonable_count': unreasonable_count,
                'analysis_coverage': (analyzed_materials / total_materials * 100) if total_materials > 0 else 0,
                'problem_rate': (problematic_materials / total_materials * 100) if total_materials > 0 else 0,
                'avg_price_variance': avg_variance,
                'estimated_savings': estimated_savings,
                'analysis_totals': analysis_totals,
                'guidance_totals': guidance_totals
            },
            'risk_stats_raw': risk_stats,
            'risk_stats': risk_distribution,
            'materials': materials,
            'analyses': analyses,
            'price_analyses': price_analyses,
            'analysis_materials_report': analysis_materials_report,
            'guidance_materials_report': guidance_materials_report,
            'top_adjustments': top_adjustments,
            'creator_name': creator_name
        }

    def _setup_document_styles(self, doc: Document):
        """统一设置文档的中文字体：标题用黑体，正文用宋体，并确保 EastAsia 字体映射"""
        try:
            normal = doc.styles['Normal']
            normal.font.name = 'SimSun'
            normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        except Exception:
            pass

        for heading in ['Heading 1', 'Heading 2', 'Heading 3']:
            try:
                h = doc.styles[heading]
                h.font.name = 'SimHei'
                h._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            except Exception:
                continue

    def _set_east_asia_font(self, run, font_name: str):
        try:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except Exception:
            pass

    def _apply_document_layout(self, doc: Document):
        try:
            for section in doc.sections:
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)
            pf = doc.styles['Normal'].paragraph_format
            pf.line_spacing = Pt(18)
            pf.space_after = Pt(6)
        except Exception:
            pass

    def _add_header_footer(self, doc: Document, report_title: str):
        try:
            for section in doc.sections:
                header_para = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
                header_para.text = report_title
                self._set_east_asia_font(header_para.runs[0], 'SimHei')
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                footer_para = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
                footer_para.text = '保密文件'
                self._set_east_asia_font(footer_para.runs[0], 'SimSun')
                footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
    
    @staticmethod
    def _resolve_region_name(code: str, map_data: Dict[str, str]) -> Optional[str]:
        if not code:
            return None
        return map_data.get(code, code)

    @staticmethod
    def _format_year_month(value: Any) -> str:
        if not value and value != 0:
            return '未设置'
        s = str(value).strip()
        if not s:
            return '未设置'
        
        # Replace '年' with '-' and remove '月'
        normalized = s.replace('年', '-').replace('月', '')
        normalized = normalized.replace('/', '-').replace('.', '-')
        
        parts = [p for p in normalized.split('-') if p]
        year = None
        month = None
        
        if len(parts) >= 2:
            year, month = parts[0], parts[1]
        elif len(normalized) == 6 and normalized.isdigit():
            year, month = normalized[:4], normalized[4:]
        elif len(normalized) == 4 and normalized.isdigit():
            year = normalized
            month = '01'
        else:
            return s
            
        try:
            year_num = int(year)
            month_num = int(month) if month else 1
            month_num = max(1, min(12, month_num))
            return f"{year_num}年{month_num:02d}月"
        except ValueError:
            return s

    def _add_title_page(self, doc: Document, data: Dict[str, Any]):
        """添加报告标题页"""
        project = data['project']
        
        # 标题
        title = doc.add_heading('造价材料审计报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 项目类型映射
        project_type_map = {
            'building': '建筑工程',
            'decoration': '装饰工程',
            'municipal': '市政工程',
            'landscape': '园林工程',
            'highway': '公路工程',
            'other': '其他工程'
        }
        
        project_type_str = project_type_map.get(project.project_type.value, '其他工程') if project.project_type else '未设置'
        
        # 地区拼接
        province_name = self._resolve_region_name(project.base_price_province, PROVINCE_MAP)
        city_name = self._resolve_region_name(project.base_price_city, CITY_MAP)
        district_name = self._resolve_region_name(project.base_price_district, DISTRICT_MAP)
        
        region_parts = []
        if province_name: region_parts.append(province_name)
        if city_name: region_parts.append(city_name)
        if district_name: region_parts.append(district_name)
        
        region = " ".join(region_parts) if region_parts else "未设置"
        
        # 工期拼接
        period = "未设置"
        start_date = self._format_year_month(project.contract_start_date)
        end_date = self._format_year_month(project.contract_end_date)
        
        if start_date != '未设置' or end_date != '未设置':
             period = f"{start_date} 至 {end_date}"
            
        # 分析范围处理
        scope_str = "未设置"
        if project.audit_scope:
            if isinstance(project.audit_scope, list):
                scope_parts = [AUDIT_SCOPE_MAP.get(s, s) for s in project.audit_scope]
                scope_str = "、".join(scope_parts)
            else:
                scope_str = str(project.audit_scope)

        # 项目信息表
        # 定义需要显示的字段
        info_items = [
            ('项目名称', project.name or ''),
            ('项目编号', project.project_code or ''),
            ('项目类型', project_type_str),
            ('项目地点', project.location or ''),
            ('业主单位', project.owner or ''),
            ('承包单位', project.contractor or ''),
            ('工程造价', f"{project.budget_amount:,.2f} 万元" if project.budget_amount else "未设置"),
            ('基期信息价日期', self._format_year_month(project.base_price_date)),
            ('合同工期', period),
            ('基期信息价地区', region),
            ('是否支持调价', "是" if project.support_price_adjustment else "否"),
            ('调价范围', f"{project.price_adjustment_range}%" if project.price_adjustment_range is not None else "未设置"),
            ('分析范围', scope_str),
            ('创建时间', project.created_at.strftime('%Y-%m-%d %H:%M:%S') if project.created_at else "未知"),
            ('创建人', data.get('creator_name', '未知')),
            ('报告日期', data['report_date'].strftime('%Y年%m月%d日')),
        ]

        table = doc.add_table(rows=len(info_items), cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, (label, value) in enumerate(info_items):
            cells = table.rows[i].cells
            cells[0].text = label
            cells[1].text = str(value)
            
            # 设置单元格对齐和字体
            for cell in cells:
                # 垂直居中
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # 段落设置
                if cell.paragraphs:
                    p = cell.paragraphs[0]
                    # 设置行间距
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    
                    # 设置字体
                    if p.runs:
                        run = p.runs[0]
                        self._set_east_asia_font(run, 'SimSun')
                        run.font.size = Pt(10.5) # 五号字
            
            # 第一列使用黑体但不再加粗，避免过重
            p = cells[0].paragraphs[0]
            if p.runs:
                run = p.runs[0]
                run.bold = False
                self._set_east_asia_font(run, 'SimHei')

        # 项目描述 (如果有)
        if project.description:
            doc.add_paragraph() # 空行
            desc_title = doc.add_paragraph()
            desc_title.add_run("项目描述：").bold = True
            self._set_east_asia_font(desc_title.runs[0], 'SimHei')
            
            desc_content = doc.add_paragraph(project.description)
            if desc_content.runs:
                self._set_east_asia_font(desc_content.runs[0], 'SimSun')
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, data: Dict[str, Any]):
        """添加执行摘要"""
        doc.add_heading('执行摘要', 1)
        
        stats = data['statistics']
        top_adjustments = data.get('top_adjustments', [])

        overview_para = doc.add_paragraph()
        run = overview_para.add_run('总体结论：')
        run.bold = True
        self._set_east_asia_font(run, 'SimHei')
        run.font.size = Pt(11)
        overview_para.add_run(
            f" 本次对项目共审查 {stats['total_materials']} 项材料，其中 {stats['analyzed_materials']} 项完成价格分析，"
            f"分析覆盖率 {stats['analysis_coverage']:.1f}%。检测到 {stats['problematic_materials']} 项存在风险，"
            f"问题发现率 {stats['problem_rate']:.1f}%，其中 {stats['unreasonable_count']} 项价格被判定为不合理。"
        )

        highlight_para = doc.add_paragraph()
        highlight_para.add_run('关键指标：').bold = True
        metrics = doc.add_table(rows=2, cols=3)
        metrics.style = 'Table Grid'
        metric_items = [
            ('平均价格偏差', f"{stats['avg_price_variance']:.2f}%"),
            ('预估节约金额', f"{stats['estimated_savings']:.2f} 元"),
            ('送审总额（无信息价）', f"{stats['analysis_totals']['original_total']:,.2f} 元")
        ]
        metric_items.extend([
            ('AI核审总额（无信息价）', f"{stats['analysis_totals']['ai_total']:,.2f} 元"),
            ('送审总额（市场信息价）', f"{stats['guidance_totals']['original_total']:,.2f} 元"),
            ('AI核审总额（市场信息价）', f"{stats['guidance_totals']['ai_total']:,.2f} 元")
        ])
        for idx, (label, value) in enumerate(metric_items):
            row = metrics.rows[idx // 3]
            cell = row.cells[idx % 3]
            cell.text = f"{label}\n{value}"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = cell.paragraphs[0].runs[0]
            self._set_east_asia_font(run, 'SimSun')
            run.font.size = Pt(9)

        doc.add_paragraph()

        if top_adjustments:
            doc.add_paragraph('重点关注材料（按核增减额排序）：').runs[0].bold = True
            for name, value in top_adjustments:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(f"{name}：核增减额 {value:+,.2f} 元")

        doc.add_paragraph(
            '建议优先对高风险及金额差异较大的材料进行复核和询价，必要时补充市场调查或合同条款校验，确保成本控制在合理范围。'
        )

        doc.add_page_break()
    
    def _add_project_overview(self, doc: Document, data: Dict[str, Any]):
        """添加项目概况"""
        doc.add_heading('项目概况', 1)
        
        project = data['project']
        stats = data['statistics']
        
        # 项目基本信息
        doc.add_heading('基本信息', 2)
        
        info_text = f"""
项目名称：{project.name}
项目编号：{project.project_code or '未提供'}
项目地点：{project.location or '未提供'}
业主单位：{project.owner or '未提供'}
承包单位：{project.contractor or '未提供'}
项目状态：{project.status.value if project.status else '未知'}
创建时间：{project.created_at.strftime('%Y年%m月%d日') if project.created_at else '未知'}
        """
        
        doc.add_paragraph(info_text.strip())
        
        # 材料统计
        doc.add_heading('材料统计', 2)
        
        stats_table = doc.add_table(rows=5, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ('材料总数', f"{stats['total_materials']}项"),
            ('已分析材料', f"{stats['analyzed_materials']}项"),
            ('问题材料', f"{stats['problematic_materials']}项"),
            ('分析覆盖率', f"{stats['analysis_coverage']:.1f}%"),
            ('问题发现率', f"{stats['problem_rate']:.1f}%")
        ]
        
        for i, (label, value) in enumerate(stats_data):
            cells = stats_table.rows[i].cells
            cells[0].text = label
            cells[1].text = value

    def _add_methodology_section(self, doc: Document, data: Dict[str, Any]):
        """添加分析方法与数据来源说明"""
        doc.add_heading('分析方法与数据来源', 1)

        method_para = doc.add_paragraph()
        method_para.add_run('分析流程：').bold = True
        method_para.add_run(
            ' 1) 基于项目材料清单提取无信息价材料与匹配的市场信息价材料；'
            ' 2) 对无信息价材料调用 AI 价格分析模型，生成推荐单价区间及合理性判断；'
            ' 3) 对市场信息价材料按照区县→市→省三级层次匹配最新信息价；'
            ' 4) 比对送审价格与 AI/市场价格，形成核增减额、风险等级和建议。'
        )

        data_para = doc.add_paragraph()
        data_para.add_run('数据来源：').bold = True
        data_para.add_run(
            ' 项目材料数据来自系统上传的清单；信息价数据引用平台最新刊期；'
            ' AI 模型训练及单位换算遵循系统既定规则，自动记录分析时间与参数。'
        )

        assumptions = doc.add_paragraph()
        assumptions.add_run('校验说明：').bold = True
        assumptions.add_run(
            ' 报告生成时自动校验计量单位、数量、价格区间，若存在缺失或异常数据，将在问题材料章节提示人工复核。'
        )

        doc.add_paragraph()

    async def _add_analysis_results(self, doc: Document, data: Dict[str, Any], include_details: bool = True, chart_files: Optional[List[Union[bytes, BytesIO, str]]] = None):
        """添加分析结果"""
        doc.add_heading('分析结果', 1)
        
        project = data['project']
        analysis_materials = data.get('analysis_materials_report', [])
        guidance_materials = data.get('guidance_materials_report', [])
        stats = data.get('statistics', {})

        doc.add_paragraph(
            f"本节汇总无信息价材料与市场信息价材料的价格对比情况，并列出存在较大差异或风险的材料明细。"
        )

        # 差额概况
        analysis_totals = stats.get('analysis_totals', {})
        guidance_totals = stats.get('guidance_totals', {})

        overview_table = doc.add_table(rows=3, cols=4)
        overview_table.style = 'Table Grid'
        headers = ['类别', '送审总额（元）', 'AI核审总额（元）', '核增减额（元）']
        for idx, header in enumerate(headers):
            cell = overview_table.rows[0].cells[idx]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'SimHei'
            run.font.size = Pt(9)
            run.bold = True

        overview_rows = [
            ('无信息价材料', analysis_totals.get('original_total', 0), analysis_totals.get('ai_total', 0), analysis_totals.get('adjustment_total', 0)),
            ('市场信息价材料', guidance_totals.get('original_total', 0), guidance_totals.get('ai_total', 0), guidance_totals.get('adjustment_total', 0))
        ]

        for row_idx, row_data in enumerate(overview_rows, start=1):
            row = overview_table.rows[row_idx]
            row.cells[0].text = row_data[0]
            row.cells[1].text = f"{row_data[1]:,.2f}"
            row.cells[2].text = f"{row_data[2]:,.2f}"
            row.cells[3].text = f"{row_data[3]:+,.2f}"
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.name = 'SimSun'
                run.font.size = Pt(9)

        doc.add_paragraph()

        # 插入图表
        if chart_files:
            # 使用表格来布局图片，这样更稳定
            # 创建一个单列的表格，每行放一张图片
            chart_table = doc.add_table(rows=len(chart_files), cols=1)
            chart_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            chart_table.autofit = False 
            chart_table.allow_autofit = False
            
            for i, chart_data in enumerate(chart_files):
                try:
                    cell = chart_table.rows[i].cells[0]
                    # 设置单元格宽度
                    cell.width = Inches(6.5)
                    
                    # 在单元格中添加段落并插入图片
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    
                    # 处理不同类型的图片数据，统一转换为bytes，确保WPS兼容性
                    image_data = None
                    if isinstance(chart_data, bytes):
                        # 直接使用字节数据
                        image_data = chart_data
                    elif isinstance(chart_data, BytesIO):
                        # 从BytesIO读取所有数据到内存，确保数据完整
                        chart_data.seek(0)
                        image_data = chart_data.read()
                        chart_data.seek(0)  # 恢复原位置
                    elif isinstance(chart_data, str) and os.path.exists(chart_data):
                        # 从文件读取内容，确保图片数据被完整读取
                        try:
                            with open(chart_data, 'rb') as f:
                                image_data = f.read()
                            logger.debug(f"成功从文件读取图片，大小: {len(image_data)} 字节")
                        except Exception as e:
                            logger.error(f"读取图片文件失败 {chart_data}: {e}")
                            continue
                    else:
                        logger.warning(f"图表数据格式不正确或文件不存在: {chart_data}")
                        continue
                    
                    # 确保图片数据存在且有效后再插入
                    if image_data and len(image_data) > 0:
                        try:
                            # 验证是否为有效的PNG图片（检查PNG文件头）
                            if image_data[:8] != b'\x89PNG\r\n\x1a\n':
                                logger.warning(f"图片数据可能不是有效的PNG格式，但仍尝试插入")
                            
                            # 创建新的BytesIO对象，确保每次插入时数据都是完整的
                            image_bytes = BytesIO(image_data)
                            
                            # 插入图片，python-docx会自动处理嵌入到Word文档
                            # 使用width参数确保图片大小合适
                            # 注意：为了WPS兼容性，确保图片是inline嵌入而不是floating
                            inline_shape = run.add_picture(image_bytes, width=Inches(6.2))
                            
                            # 记录成功信息（添加WPS兼容性标记）
                            logger.info(f"[WPS兼容性修复] 图片已成功嵌入Word文档，大小: {len(image_data)} 字节，类型: {type(inline_shape)}")
                            
                            # 确保图片是inline方式嵌入（WPS兼容性要求）
                            # python-docx的add_picture默认就是inline方式，但我们可以验证
                            try:
                                # 验证图片确实被嵌入
                                if hasattr(inline_shape, '_inline'):
                                    logger.debug("图片已确认为inline嵌入方式")
                            except Exception as e:
                                logger.warning(f"无法验证图片嵌入方式: {e}，但图片已插入")
                            
                        except Exception as e:
                            logger.error(f"插入图片到Word文档失败: {e}", exc_info=True)
                    else:
                        logger.warning(f"图片数据为空，跳过插入")
                        
                    # 确保没有多余的边框（如果不需要边框的话）
                    # self._remove_table_borders(chart_table) 
                    
                except Exception as e:
                    logger.error(f"添加图表失败: {e}", exc_info=True)
            
            doc.add_paragraph()  # 表格后添加空行分开

        if include_details:
            # 生成表1：材料价格分析表（无信息价材料）
            if analysis_materials:
                # 筛选核增（减）额为正数的数据（即送审价 > AI价，核减）
                filtered_materials = [
                    m for m in analysis_materials
                    if m.get('adjustment', 0) > 0
                ]
                # 按权重百分比降序排序（权重高的材料排在前面）
                filtered_materials.sort(
                    key=lambda x: x.get('weightPercentage', 0),
                    reverse=True
                )
                # 下载的报告仅展示权重最高的前15条记录
                filtered_materials = filtered_materials[:15]

                if filtered_materials:
                    self.create_standard_material_table(
                        doc, 
                        "表1 材料价格分析表（无信息价材料）", 
                        filtered_materials,
                        "analysis",
                        project.name or "未命名项目"
                    )
                else:
                    p = doc.add_paragraph("表1：未发现正向核增（减）材料（无核减项）。")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 生成表2：材料价格分析表（市场信息价材料）
            if guidance_materials:
                self.create_standard_material_table(
                    doc, 
                    "表2 材料价格分析表（市场信息价材料）", 
                    guidance_materials,  # 显示所有数据
                    "guidance_price",
                    project.name or "未命名项目"
                )
            elif not guidance_materials:
                # 如果没有市场信息价材料，显示提示信息
                doc.add_paragraph("暂无市场信息价材料数据")
        else:
            doc.add_paragraph("（详细清单已省略，请查看摘要或选择生成完整报告）")
        
        # 风险等级分布摘要
        doc.add_heading('风险等级分布', 2)
        
        risk_stats = data.get('risk_stats', {})
        if risk_stats:
            risk_table = doc.add_table(rows=len(risk_stats) + 1, cols=2)
            risk_table.style = 'Table Grid'
            
            # 表头
            cells = risk_table.rows[0].cells
            cells[0].text = '风险等级'
            cells[1].text = '材料数量'
            for c in cells:
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].bold = True
                self._set_east_asia_font(p.runs[0], 'SimHei')
            
            # 数据行
            for i, (risk_level, count) in enumerate(risk_stats.items(), 1):
                cells = risk_table.rows[i].cells
                cells[0].text = risk_level
                cells[1].text = f"{count}项"
        
        # 价格分析汇总
        doc.add_heading('价格分析汇总', 2)
        
        stats = data.get('statistics', {})
        summary_text = f"""
• 平均价格偏差：{stats.get('avg_price_variance', 0):.2f}%
• 不合理价格材料：{stats.get('unreasonable_count', 0)}项
• 预估节约金额：{stats.get('estimated_savings', 0):.2f}元
        """
        
        doc.add_paragraph(summary_text.strip())
    
    def _add_problematic_materials(self, doc: Document, data: Dict[str, Any]):
        """添加问题材料详情"""
        doc.add_heading('问题材料详情', 1)
        
        materials = data['materials']
        price_analyses = data['price_analyses']
        
        # 筛选问题材料
        problematic_materials = [m for m in materials if m.is_problematic]
        
        if not problematic_materials:
            doc.add_paragraph('未发现问题材料。')
            return
        
        # 创建问题材料表格
        table = doc.add_table(rows=len(problematic_materials) + 1, cols=6)
        table.style = 'Table Grid'
        
        # 表头
        headers = ['序号', '材料名称', '规格', '单价', '风险等级', '问题描述']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            self._set_east_asia_font(p.runs[0], 'SimHei')
        
        # 数据行
        for i, material in enumerate(problematic_materials, 1):
            cells = table.rows[i].cells
            cells[0].text = str(i)
            cells[1].text = material.material_name
            cells[2].text = material.specification or ''
            cells[3].text = f"{material.unit_price:.2f}" if material.unit_price else ''
            
            # 获取风险等级和问题描述
            if material.id in price_analyses:
                analysis = price_analyses[material.id]
                cells[4].text = analysis.risk_level or ''
                
                # 生成问题描述
                problems = []
                if analysis.is_reasonable == False:
                    problems.append('价格不合理')
                if analysis.price_variance and abs(analysis.price_variance) > 30:
                    problems.append(f'价格偏差{analysis.price_variance:.1f}%')
                
                cells[5].text = '；'.join(problems) or '需要人工确认'
            else:
                cells[4].text = '未分析'
                cells[5].text = '缺少价格分析'
    
    def _process_chart_images(self, chart_images: Dict[str, str]) -> List[bytes]:
        """处理前端传入的图表图片，返回图片字节数据列表以便直接嵌入Word文档
        返回bytes而不是BytesIO，确保数据完整性，提高WPS兼容性
        """
        chart_data_list = []
        
        # 定义图表顺序，与前端对应
        # risk: 风险等级分布
        # adjustment: 核增减额TOP10
        # totals: 送审 VS AI 核审 总额对比
        order = ['risk', 'adjustment', 'totals']
        
        for key in order:
            base64_str = chart_images.get(key)
            if not base64_str:
                continue
                
            try:
                # 移除base64前缀
                if ',' in base64_str:
                    base64_str = base64_str.split(',')[1]
                    
                image_data = base64.b64decode(base64_str)
                
                # 直接保存字节数据，而不是BytesIO对象，确保数据完整性
                chart_data_list.append(image_data)
                
                logger.info(f"[WPS兼容性修复] 成功处理图表 {key}，数据大小: {len(image_data)} 字节，PNG头: {image_data[:8].hex() if len(image_data) >= 8 else 'N/A'}")
            except Exception as e:
                logger.error(f"处理图表 {key} 失败: {e}")
                
        return chart_data_list

    def _generate_charts(self, data: Dict[str, Any]) -> List[str]:
        """生成图表文件"""
        chart_files = []
        
        try:
            chart_path = self._create_risk_level_pie_chart(data)
            if chart_path:
                chart_files.append(chart_path)
            
            chart_path = self._create_adjustment_top_chart(data)
            if chart_path:
                chart_files.append(chart_path)

            chart_path = self._create_total_comparison_chart(data)
            if chart_path:
                chart_files.append(chart_path)

            # 新增：价格偏差分析图（如果有分析结果）
            chart_path = self._create_price_variance_chart(data)
            if chart_path:
                chart_files.append(chart_path)

            # 如果以上图表都没有生成（例如项目数据较少），生成一张占位说明图
            if not chart_files:
                chart_path = self._create_no_data_chart()
                if chart_path:
                    chart_files.append(chart_path)

        except Exception as e:
            logger.error(f"生成图表失败: {e}")
        
        return chart_files
    
    def _create_risk_level_pie_chart(self, data: Dict[str, Any]) -> Optional[str]:
        """创建风险等级饼图"""
        risk_stats = data.get('risk_stats_raw', {})
        if not risk_stats:
            return None
        
        # 准备数据
        label_map = {
            'normal': '正常',
            'low': '低风险',
            'medium': '中风险',
            'high': '高风险',
            'critical': '极高风险',
            'severe': '极高风险',
            'unknown': '待确认'
        }
        color_map = {
            'normal': '#67C23A',
            'low': '#409EFF',
            'medium': '#E6A23C',
            'high': '#F56C6C',
            'critical': '#C039A5',
            'severe': '#C039A5',
            'unknown': '#909399'
        }

        ordered_levels = ['normal', 'low', 'medium', 'high', 'critical', 'unknown']
        labels = []
        sizes = []
        colors = []
        for level in ordered_levels:
            if level in risk_stats:
                labels.append(label_map.get(level, level))
                sizes.append(risk_stats[level])
                colors.append(color_map.get(level, '#409EFF'))

        # 创建饼图
        plt.figure(figsize=(8, 8))
        plt.pie(sizes, labels=labels, colors=colors[:len(labels)], autopct='%1.1f%%', startangle=90)
        plt.title('材料风险等级分布', fontsize=16)
        plt.axis('equal')
        
        # 保存图表
        chart_filename = f"risk_levels_{uuid.uuid4().hex[:8]}.png"
        chart_path = self.reports_dir / chart_filename
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return str(chart_path)

    def _create_adjustment_top_chart(self, data: Dict[str, Any]) -> Optional[str]:
        items: List[Tuple[str, float]] = []
        for item in data.get('analysis_materials_report', []):
            items.append((item['materialName'], item['adjustment']))
        for item in data.get('guidance_materials_report', []):
            items.append((item['materialName'], item['adjustment']))

        if not items:
            return None

        items.sort(key=lambda x: abs(x[1]), reverse=True)
        top_items = items[:10]
        names = [name for name, _ in top_items][::-1]
        values = [value for _, value in top_items][::-1]
        colors = ['#F56C6C' if v > 0 else '#67C23A' for v in values]

        plt.figure(figsize=(12, 7))
        bars = plt.barh(names, values, color=colors)
        plt.title('核增减额TOP10', fontsize=16)
        plt.xlabel('金额（元）', fontsize=12)
        plt.grid(axis='x', alpha=0.2)
        plt.axvline(0, color='#333', linewidth=0.8)

        max_value = max((abs(v) for v in values), default=0)
        offset = max_value * 0.02 if max_value else 0
        for bar, value in zip(bars, values):
            x_pos = value + offset if value >= 0 else value - offset
            ha = 'left' if value >= 0 else 'right'
            plt.text(x_pos, bar.get_y() + bar.get_height() / 2, f"{value:+,.2f}", va='center', ha=ha, fontsize=9)

        plt.tight_layout()

        chart_filename = f"adjustment_top_{uuid.uuid4().hex[:8]}.png"
        chart_path = self.reports_dir / chart_filename
        plt.savefig(chart_path, dpi=300)
        plt.close()

        return str(chart_path)

    def _create_total_comparison_chart(self, data: Dict[str, Any]) -> Optional[str]:
        stats = data.get('statistics', {})
        analysis_totals = stats.get('analysis_totals') or {}
        guidance_totals = stats.get('guidance_totals') or {}

        categories = ['无信息价材料', '市场信息价材料']
        original_totals = [analysis_totals.get('original_total', 0), guidance_totals.get('original_total', 0)]
        ai_totals = [analysis_totals.get('ai_total', 0), guidance_totals.get('ai_total', 0)]

        if not any(original_totals) and not any(ai_totals):
            return None

        x = range(len(categories))
        width = 0.35

        plt.figure(figsize=(10, 6))
        plt.bar([i - width / 2 for i in x], original_totals, width, label='送审总额', color='#409EFF')
        plt.bar([i + width / 2 for i in x], ai_totals, width, label='AI核审总额', color='#67C23A')

        plt.xticks(list(x), categories)
        plt.ylabel('金额（元）')
        plt.title('送审 VS AI 核审 总额对比', fontsize=16)
        plt.legend()
        plt.grid(axis='y', alpha=0.2)

        for idx, value in enumerate(original_totals):
            plt.text(idx - width / 2, value, f"{value:,.2f}", ha='center', va='bottom', fontsize=9)
        for idx, value in enumerate(ai_totals):
            plt.text(idx + width / 2, value, f"{value:,.2f}", ha='center', va='bottom', fontsize=9)

        plt.tight_layout()

        chart_filename = f"totals_compare_{uuid.uuid4().hex[:8]}.png"
        chart_path = self.reports_dir / chart_filename
        plt.savefig(chart_path, dpi=300)
        plt.close()

        return str(chart_path)
    def _create_price_variance_chart(self, data: Dict[str, Any]) -> Optional[str]:
        """创建价格偏差分析图"""
        analyses = data['analyses']
        if not analyses:
            return None
        
        # 准备数据
        variances = []
        materials = []
        
        for i, analysis in enumerate(analyses[:20]):  # 只显示前20个
            if analysis.price_variance is not None:
                variances.append(analysis.price_variance)
                materials.append(f"材料{i+1}")
        
        if len(variances) < 2:
            return None
        
        # 创建条形图
        plt.figure(figsize=(12, 6))
        colors = ['red' if v > 30 else 'orange' if v > 15 else 'green' for v in variances]
        plt.bar(materials, variances, color=colors, alpha=0.7)
        plt.title('材料价格偏差分析', fontsize=16)
        plt.xlabel('材料', fontsize=12)
        plt.ylabel('价格偏差（%）', fontsize=12)
        plt.xticks(rotation=45)
        plt.axhline(y=15, color='orange', linestyle='--', alpha=0.5, label='15%警戒线')
        plt.axhline(y=30, color='red', linestyle='--', alpha=0.5, label='30%风险线')
        plt.legend()
        plt.grid(True, alpha=0.3)
        
        # 保存图表
        chart_filename = f"price_variance_{uuid.uuid4().hex[:8]}.png"
        chart_path = self.reports_dir / chart_filename
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return str(chart_path)

    def _create_no_data_chart(self) -> Optional[str]:
        """当没有可用数据时生成占位图，避免文档中图表区域完全空白"""
        try:
            plt.figure(figsize=(8, 4))
            plt.axis('off')
            plt.text(
                0.5,
                0.5,
                '当前项目暂无可展示的图表数据\n（无风险材料或差异为 0）',
                ha='center',
                va='center',
                fontsize=14
            )
            chart_filename = f"no_data_{uuid.uuid4().hex[:8]}.png"
            chart_path = self.reports_dir / chart_filename
            plt.savefig(chart_path, dpi=200, bbox_inches='tight')
            plt.close()
            return str(chart_path)
        except Exception as e:
            logger.error(f"生成占位图失败: {e}")
            return None
    
    def _add_charts_to_document(self, doc: Document, chart_files: List[str]):
        """将图表添加到文档中"""
        if not chart_files:
            return
        
        doc.add_heading('图表分析', 1)
        
        for chart_file in chart_files:
            if os.path.exists(chart_file):
                try:
                    doc.add_picture(chart_file, width=Inches(6))
                    doc.add_paragraph()  # 添加间距
                except Exception as e:
                    logger.error(f"添加图表失败 {chart_file}: {e}")
        
        doc.add_page_break()
    
    def _add_recommendations(self, doc: Document, data: Dict[str, Any]):
        """添加建议措施"""
        doc.add_heading('建议措施', 1)
        
        stats = data['statistics']
        
        recommendations = []
        
        if stats['problem_rate'] > 10:
            recommendations.append(
                "发现较多价格异常材料，建议加强供应商管理和价格监控。"
            )
        
        if stats['avg_price_variance'] > 20:
            recommendations.append(
                f"平均价格偏差达到{stats['avg_price_variance']:.1f}%，超出正常范围，"
                "建议重新核实市场价格。"
            )
        
        if stats['estimated_savings'] > 10000:
            recommendations.append(
                f"预估可节约成本{stats['estimated_savings']:.2f}元，建议优化采购策略。"
            )
        
        # 通用建议
        recommendations.extend([
            "定期更新基准材料价格数据库，确保价格信息的时效性。",
            "建立材料价格预警机制，及时发现价格异常。",
            "加强与供应商的沟通，了解价格变动原因。",
            "对高风险材料进行重点监控和审核。"
        ])
        
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}")
    
    def _add_appendices(self, doc: Document, data: Dict[str, Any]):
        """添加附录"""
        doc.add_heading('附录', 1)
        
        # 附录A: 技术说明
        doc.add_heading('附录A：技术说明', 2)
        
        tech_info = """
本报告使用人工智能技术进行材料价格分析，主要技术特点如下：

1. 多AI服务集成：支持deepseek、通义千问、豆包多个AI服务。
2. 智能匹配算法：基于材料名称、规格、单位等多维度进行相似度计算。
3. 价格合理性分析：结合统计学方法和AI预测进行综合判断。
4. 故障转移机制：确保分析服务的高可用性。

分析结果仅供参考，最终决策请结合实际情况和专业判断。
        """
        
        doc.add_paragraph(tech_info.strip())
        
        # 附录B: 数据统计
        doc.add_heading('附录B：详细统计', 2)
        
        analyses = data['analyses']
        if analyses:
            # AI服务使用统计
            ai_providers = {}
            for analysis in analyses:
                provider = analysis.analysis_model or 'unknown'
                ai_providers[provider] = ai_providers.get(provider, 0) + 1
            
            doc.add_paragraph("AI服务使用统计：")
            for provider, count in ai_providers.items():
                doc.add_paragraph(f"• {provider}: {count}次")
    
    def _set_table_borders(self, table):
        """设置表格边框"""
        tbl = table._tbl
        for cell in table._cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # 创建边框元素
            tcBorders = OxmlElement('w:tcBorders')
            
            # 设置所有边框
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')  # 边框粗细
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # 黑色边框
                tcBorders.append(border)
            
            tcPr.append(tcBorders)
    
    def _merge_cells_horizontally(self, table, row_idx, start_col, end_col):
        """水平合并单元格"""
        if start_col >= end_col:
            return
        
        # 获取要合并的单元格
        cells = table.rows[row_idx].cells
        main_cell = cells[start_col]
        
        # 合并后续单元格
        for col_idx in range(start_col + 1, end_col + 1):
            if col_idx < len(cells):
                main_cell.merge(cells[col_idx])
    
    def create_standard_material_table(self, doc: Document, title: str, materials_data: List[Dict], 
                                     table_type: str = "analysis", project_name: str = "") -> None:
        """创建标准化的材料分析表格
        
        Args:
            doc: Word文档对象
            title: 表格标题
            materials_data: 材料数据列表
            table_type: 表格类型 ("analysis" 或 "guidance_price")
            project_name: 项目名称
        """
        # 插入新的分节符，设置为横向
        if len(doc.sections) > 0:
            doc.add_section(WD_SECTION.NEW_PAGE)
        
        # 设置当前节为横向
        current_section = doc.sections[-1]
        current_section.orientation = WD_ORIENT.LANDSCAPE
        
        # 设置横向页面的页边距（交换宽度和高度的概念）
        current_section.page_height = Cm(21)    # A4纸宽度
        current_section.page_width = Cm(29.7)   # A4纸长度
        current_section.left_margin = Cm(1.5)
        current_section.right_margin = Cm(1.5)
        current_section.top_margin = Cm(1.5)
        current_section.bottom_margin = Cm(1.5)
        
        # 添加表格标题
        heading = doc.add_heading(title, level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置标题字体为黑体
        for run in heading.runs:
            run.font.name = 'SimHei'  # 黑体
            run.font.size = Pt(14)
            run.bold = True
        
        # 添加项目信息和单位信息在同一行
        project_info = doc.add_paragraph()
        
        # 左侧项目名称
        run1 = project_info.add_run("项目名称：")
        run1.bold = True
        run1.font.name = 'SimSun'  # 宋体
        run1.font.size = Pt(10)
        
        run2 = project_info.add_run(project_name or "未命名项目")
        run2.font.name = 'SimSun'
        run2.font.size = Pt(10)
        
        # 添加制表符来推送"单位：元"到右侧
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
        
        # 设置制表位到页面右侧
        tab_stops = project_info.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(25), WD_TAB_ALIGNMENT.RIGHT)  # 25cm处右对齐制表位
        
        # 添加制表符和单位信息
        run3 = project_info.add_run("\t单位：元")
        run3.bold = True
        run3.font.name = 'SimSun'
        run3.font.size = Pt(10)
        
        # 根据表类型分别构建列结构
        if table_type == "guidance_price":
            # ==================== 市场信息价材料表结构 ====================
            # 列：序号、材料名称、规格型号、单位、数量、
            #     送审结算(单价/合价)、市场信息价结算(基期信息价/价格差异/合同期平均价/风险幅度)、
            #     调差、权重（%）
            base_cols = 13
            table = doc.add_table(rows=2, cols=base_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml

            tbl = table._tbl
            tblPr = tbl.tblPr
            tblBorders = parse_xml(
                r'<w:tblBorders {}>'
                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'</w:tblBorders>'.format(nsdecls('w'))
            )
            tblPr.append(tblBorders)

            widths = [
                Cm(1.0),  # 序号
                Cm(3.0),  # 材料名称
                Cm(2.5),  # 规格型号
                Cm(1.0),  # 单位
                Cm(1.5),  # 数量
                Cm(1.8),  # 送审单价
                Cm(1.8),  # 送审合价
                Cm(1.8),  # 基期信息价
                Cm(1.8),  # 价格差异
                Cm(1.8),  # 合同期平均价
                Cm(1.8),  # 风险幅度
                Cm(2.0),  # 调差
                Cm(1.5),  # 权重
            ]
            for i, width in enumerate(widths[:base_cols]):
                for row in table.rows:
                    if i < len(row.cells):
                        row.cells[i].width = width

            # 第一行表头
            header_row1 = table.rows[0]
            headers1 = [
                '序号', '材料名称', '规格型号', '单位', '数量',
                '送审结算', '',
                '市场信息价结算', '', '', '',
                '调差', '权重（%）'
            ]
            for i, header in enumerate(headers1[:base_cols]):
                if header:
                    cell = header_row1.cells[i]
                    cell.text = header
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = 'SimHei'
                    run.font.size = Pt(9)
                    run.bold = True

            # 合并第一行多级表头
            self._merge_cells_horizontally(table, 0, 5, 6)   # 送审结算
            self._merge_cells_horizontally(table, 0, 7, 10)  # 市场信息价结算

            # 第二行表头
            header_row2 = table.rows[1]
            headers2 = [
                '', '', '', '', '',
                '单价', '合价',
                '基期信息价', '价格差异', '合同期平均价', '风险幅度',
                '', ''
            ]
            for i, header in enumerate(headers2[:base_cols]):
                if header:
                    cell = header_row2.cells[i]
                    cell.text = header
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = 'SimHei'
                    run.font.size = Pt(8)
                    run.bold = True

            # 合并无需子表头的列（首行和次行）
            cols_to_merge = [0, 1, 2, 3, 4, 11, 12]
            for col_idx in cols_to_merge:
                if col_idx < base_cols:
                    try:
                        table.cell(0, col_idx).merge(table.cell(1, col_idx))
                    except Exception:
                        pass

            # 数据行
            for idx, material in enumerate(materials_data):
                row = table.add_row()
                data_cells = [
                    str(idx + 1),
                    material.get('materialName', ''),
                    material.get('specification', ''),
                    material.get('unit', ''),
                    f"{material.get('quantity', 0):,.0f}",
                    f"{material.get('originalUnitPrice', 0):,.4f}",
                    f"{material.get('originalTotalPrice', 0):,.2f}",
                    f"{material.get('originalBasePrice', 0):,.4f}",
                    f"{material.get('priceDiff', 0):+,.4f}",
                    f"{material.get('aiUnitPrice', 0):,.4f}",
                    f"{material.get('riskRate', 0):.2f}%",
                    f"{material.get('adjustment', 0):+,.2f}",
                    f"{material.get('weightPercentage', 0):.2f}",
                ]
                for j, data in enumerate(data_cells[:base_cols]):
                    cell = row.cells[j]
                    cell.text = data
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = 'SimSun'
                    run.font.size = Pt(8)

            # 合计行
            summary_row = table.add_row()
            summary_data = ['',
                            '合计', '', '', '',
                            '', '', '', '', '',
                            '',
                            '', '']

            total_original = sum(m.get('originalTotalPrice', 0) for m in materials_data)
            total_guidance = sum(m.get('aiTotalPrice', 0) for m in materials_data)
            total_adjustment = sum(m.get('adjustment', 0) for m in materials_data)
            total_weight = sum(m.get('weightPercentage', 0) for m in materials_data)

            summary_data[6] = f"{total_original:,.2f}"       # 送审合价合计
            summary_data[9] = f"{total_guidance:,.2f}"       # 合同期平均价合计
            summary_data[11] = f"{total_adjustment:+,.2f}"   # 调差合计
            summary_data[12] = f"{total_weight:.2f}"         # 权重合计

            for j, data in enumerate(summary_data[:base_cols]):
                cell = summary_row.cells[j]
                cell.text = data
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.name = 'SimSun'
                run.font.size = Pt(8)
                run.bold = True

        else:
            # ==================== 无信息价材料表结构（原逻辑） ====================
            base_cols = 11  # 序号、材料名称、规格型号、单位、数量、送审结算(单价/合价)、AI辅助审核(单价/合价)、核增减额、权重
            table = doc.add_table(rows=2, cols=base_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml

            tbl = table._tbl
            tblPr = tbl.tblPr
            tblBorders = parse_xml(
                r'<w:tblBorders {}>'
                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'</w:tblBorders>'.format(nsdecls('w'))
            )
            tblPr.append(tblBorders)

            widths = [Cm(1.0), Cm(3.0), Cm(2.5), Cm(1.0), Cm(1.5),
                      Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8), Cm(2.0), Cm(1.5)]
            for i, width in enumerate(widths[:base_cols]):
                for row in table.rows:
                    if i < len(row.cells):
                        row.cells[i].width = width

            header_row1 = table.rows[0]
            headers1 = ['序号', '材料名称', '规格型号', '单位', '数量',
                        '送审结算', '', 'AI 辅助审核', '', '核增（减）额', '权重（%）']
            for i, header in enumerate(headers1[:base_cols]):
                if header:
                    cell = header_row1.cells[i]
                    cell.text = header
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = 'SimHei'
                    run.font.size = Pt(9)
                    run.bold = True

            self._merge_cells_horizontally(table, 0, 5, 6)
            self._merge_cells_horizontally(table, 0, 7, 8)

            header_row2 = table.rows[1]
            headers2 = ['', '', '', '', '', '单价', '合价', '单价', '合价', '', '']
            for i, header in enumerate(headers2[:base_cols]):
                if header:
                    cell = header_row2.cells[i]
                    cell.text = header
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = 'SimHei'
                    run.font.size = Pt(8)
                    run.bold = True

            cols_to_merge = [0, 1, 2, 3, 4, 9, 10]
            for col_idx in cols_to_merge:
                if col_idx < base_cols:
                    try:
                        table.cell(0, col_idx).merge(table.cell(1, col_idx))
                    except Exception:
                        pass

            for idx, material in enumerate(materials_data):
                row = table.add_row()
                data_cells = [
                    str(idx + 1),
                    material.get('materialName', ''),
                    material.get('specification', ''),
                    material.get('unit', ''),
                    f"{material.get('quantity', 0):,.0f}",
                    f"{material.get('originalUnitPrice', 0):,.2f}",
                    f"{material.get('originalTotalPrice', 0):,.2f}",
                    f"{material.get('aiUnitPrice', 0):,.2f}",
                    f"{material.get('aiTotalPrice', 0):,.2f}",
                    f"{material.get('adjustment', 0):+,.2f}",
                    f"{material.get('weightPercentage', 0):.2f}",
                ]
                for j, data in enumerate(data_cells[:base_cols]):
                    cell = row.cells[j]
                    cell.text = data
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = 'SimSun'
                    run.font.size = Pt(8)

            summary_row = table.add_row()
            summary_data = ['', '合计', '', '', '', '', '', '', '', '', '']

            total_original = sum(m.get('originalTotalPrice', 0) for m in materials_data)
            total_ai = sum(m.get('aiTotalPrice', 0) for m in materials_data)
            total_adjustment = sum(m.get('adjustment', 0) for m in materials_data)
            total_weight = sum(m.get('weightPercentage', 0) for m in materials_data)

            summary_data[6] = f"{total_original:,.2f}"
            summary_data[8] = f"{total_ai:,.2f}"
            summary_data[9] = f"{total_adjustment:+,.2f}"
            summary_data[10] = f"{total_weight:.2f}"

            for j, data in enumerate(summary_data[:base_cols]):
                cell = summary_row.cells[j]
                cell.text = data
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.name = 'SimSun'
                run.font.size = Pt(8)
                run.bold = True
        
        # 添加备注
        notes = doc.add_paragraph()
        if table_type == "analysis":
            notes.text = "备注：（1）本表可扩展；（2）差额为正值即核减，负值即核增；（3）本表可作为过程资料一并归档。"
        else:
            notes.text = "备注：（1）本表可扩展；（2）差额为正值即核减，负值即核增；（3）本表可纳入审价过程资料一并归档。"
        
        notes.runs[0].font.name = 'SimSun'  # 宋体
        notes.runs[0].font.size = Pt(8)
        
        doc.add_paragraph()  # 添加间距
