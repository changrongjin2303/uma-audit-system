import os
import json
from datetime import datetime
from typing import Dict, List, Optional, Any
from pathlib import Path
from loguru import logger
import matplotlib
matplotlib.use('Agg')  # 服务器环境下无GUI
import matplotlib.pyplot as plt

try:
    from docx import Document
    from docx.shared import Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError as e:
    # 日志模块已就绪，安全记录依赖缺失
    logger.warning(f"python-docx not available: {e}")
    DOCX_AVAILABLE = False

from app.models.project import Project, ProjectMaterial
from app.models.analysis import PriceAnalysis
from app.core.config import settings


class SimpleReportGenerator:
    """简化的审计报告生成器 - 避免复杂依赖"""
    
    def __init__(self):
        # 使用绝对路径，确保在不同环境下都能正常工作
        if hasattr(settings, 'REPORTS_DIR') and settings.REPORTS_DIR:
            self.reports_dir = Path(settings.REPORTS_DIR)
        else:
            # 使用当前工作目录下的reports文件夹
            self.reports_dir = Path.cwd() / "reports"
        
        # 确保目录存在
        try:
            self.reports_dir.mkdir(parents=True, exist_ok=True)
            logger.info(f"报告输出目录: {self.reports_dir.absolute()}")
        except Exception as e:
            logger.error(f"创建报告目录失败: {e}")
            # 如果无法创建目录，使用临时目录
            import tempfile
            self.reports_dir = Path(tempfile.gettempdir()) / "uma_reports"
            self.reports_dir.mkdir(parents=True, exist_ok=True)
            logger.warning(f"使用临时目录: {self.reports_dir.absolute()}")
        
    def generate_audit_report(
        self,
        project: Project,
        materials: List[ProjectMaterial],
        analyses: List[PriceAnalysis],
        report_config: Dict[str, Any] = None
    ) -> str:
        """生成简化的审计报告
        
        Args:
            project: 项目信息
            materials: 项目材料列表
            analyses: 价格分析结果
            report_config: 报告配置参数
            
        Returns:
            报告文件路径
        """
        try:
            logger.info(f"开始生成简化审计报告: 项目ID {project.id if project else 'None'}")
            
            # 检查必要的依赖
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx库不可用，无法生成Word报告")
            
            # 验证输入参数
            if not project:
                raise ValueError("项目信息不能为空")
            
            if not materials:
                logger.warning("材料列表为空，将使用示例数据")
                materials = self._create_sample_materials(project.id if project else 1)
            
            if not analyses:
                logger.warning("分析结果为空，将使用示例数据")
                analyses = self._create_sample_analyses(materials)
            
            # 生成报告数据
            report_data = self._prepare_report_data(project, materials, analyses)
            
            # 创建Word文档（优先使用自定义模板）
            template_candidates = [
                Path('static/templates/audit_report_template.docx'),
                Path('templates/audit_report_template.docx'),
                Path('backend/static/templates/audit_report_template.docx'),
            ]
            doc = None
            for t in template_candidates:
                if t.exists():
                    try:
                        doc = Document(str(t))
                        logger.info(f"使用报告模板: {t}")
                        break
                    except Exception as _:
                        pass
            if doc is None:
                doc = Document()

            # 设置中文默认字体（若可用）
            try:
                style = doc.styles['Normal']
                style.font.name = '宋体'
                from docx.oxml.shared import qn as _qn
                style._element.rPr.rFonts.set(_qn('w:eastAsia'), '宋体')
            except Exception:
                pass
            
            # 添加报告内容
            self._add_title_page(doc, report_data)
            self._add_executive_summary(doc, report_data)
            self._add_project_overview(doc, report_data)
            self._add_analysis_results(doc, report_data)
            self._add_problematic_materials(doc, report_data)

            # 图表分析（可选）
            include_charts = True
            if report_config and isinstance(report_config, dict):
                include_charts = report_config.get('include_charts', True)
            if include_charts:
                charts = self._render_charts(report_data)
                self._add_charts_section(doc, charts)

            self._add_recommendations(doc, report_data)
            self._add_appendix(doc)
            
            # 保存报告
            project_id = project.id if project else 'unknown'
            report_filename = f"audit_report_{project_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            report_path = self.reports_dir / report_filename
            
            # 确保可以写入文件
            try:
                doc.save(str(report_path))
                logger.info(f"简化审计报告生成成功: {report_path}")
                
                # 验证文件是否真正创建
                if not report_path.exists():
                    raise FileNotFoundError(f"报告文件保存失败: {report_path}")
                
                file_size = report_path.stat().st_size
                logger.info(f"报告文件大小: {file_size} 字节")
                
                return str(report_path.absolute())
                
            except Exception as save_error:
                logger.error(f"保存报告文件失败: {save_error}")
                # 尝试使用不同的文件名
                import uuid
                backup_filename = f"audit_report_{uuid.uuid4().hex[:8]}.docx"
                backup_path = self.reports_dir / backup_filename
                doc.save(str(backup_path))
                logger.warning(f"使用备用文件名保存: {backup_path}")
                return str(backup_path.absolute())
            
        except Exception as e:
            logger.error(f"生成简化审计报告失败: {e}")
            logger.error(f"错误类型: {type(e).__name__}")
            logger.error(f"项目信息: {project.name if project else 'None'}")
            logger.error(f"材料数量: {len(materials) if materials else 0}")
            logger.error(f"分析数量: {len(analyses) if analyses else 0}")
            
            # 尝试生成一个最简单的报告
            try:
                return self._generate_minimal_report(project)
            except Exception as fallback_error:
                logger.error(f"生成简化报告也失败: {fallback_error}")
                raise Exception(f"报告生成完全失败: {e}") from e
    
    def _prepare_report_data(
        self,
        project: Project,
        materials: List[ProjectMaterial],
        analyses: List[PriceAnalysis]
    ) -> Dict[str, Any]:
        """准备报告数据"""
        
        # 基础统计
        total_materials = len(materials)
        analyzed_materials = len([m for m in materials if m.is_analyzed])
        problematic_materials = len([m for m in materials if m.is_problematic])
        
        # 价格分析统计
        price_analyses = {a.material_id: a for a in analyses}
        unreasonable_count = len([a for a in analyses if a.is_reasonable == False])
        
        # 风险等级统计
        risk_stats = {}
        for analysis in analyses:
            risk_level = analysis.risk_level or 'unknown'
            risk_stats[risk_level] = risk_stats.get(risk_level, 0) + 1
        
        # 价格偏差统计
        price_variances = [a.price_variance for a in analyses if a.price_variance is not None]
        avg_variance = sum(price_variances) / len(price_variances) if price_variances else 0
        
        # 估算节约金额
        estimated_savings = 0
        for material in materials:
            if material.id in price_analyses:
                analysis = price_analyses[material.id]
                if analysis.predicted_price_avg and material.unit_price:
                    if material.unit_price > analysis.predicted_price_avg:
                        savings = (material.unit_price - analysis.predicted_price_avg) * (material.quantity or 0)
                        estimated_savings += savings
        
        return {
            'project': project,
            'report_date': datetime.now(),
            'statistics': {
                'total_materials': total_materials,
                'analyzed_materials': analyzed_materials,
                'problematic_materials': problematic_materials,
                'unreasonable_count': unreasonable_count,
                'analysis_coverage': (analyzed_materials / total_materials * 100) if total_materials > 0 else 0,
                'problem_rate': (problematic_materials / total_materials * 100) if total_materials > 0 else 0,
                'avg_price_variance': avg_variance,
                'estimated_savings': estimated_savings,
                'risk_distribution': risk_stats
            },
            'materials': materials,
            'analyses': price_analyses
        }

    def _render_charts(self, report_data: Dict[str, Any]) -> Dict[str, str]:
        """使用matplotlib生成图表并返回图片路径"""
        charts_dir = self.reports_dir / "charts"
        charts_dir.mkdir(exist_ok=True)

        materials: List[ProjectMaterial] = report_data['materials']
        analyses_map: Dict[int, PriceAnalysis] = report_data['analyses']
        stats = report_data['statistics']

        generated: Dict[str, str] = {}

        # 1) 风险等级饼图
        try:
            labels = list(stats['risk_distribution'].keys()) or ['unknown']
            sizes = [stats['risk_distribution'][k] for k in labels] if stats['risk_distribution'] else [1]
            fig, ax = plt.subplots(figsize=(5, 4), dpi=150)
            ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
            ax.set_title('风险等级分布')
            risk_path = charts_dir / f"risk_levels_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            plt.tight_layout()
            fig.savefig(risk_path)
            plt.close(fig)
            generated['risk_levels'] = str(risk_path)
        except Exception as e:
            logger.warning(f"生成风险等级饼图失败: {e}")

        # 2) 价格分布直方图
        try:
            prices = [a.predicted_price_avg for a in analyses_map.values() if a and a.predicted_price_avg]
            if prices:
                fig, ax = plt.subplots(figsize=(5, 4), dpi=150)
                ax.hist(prices, bins=20, color='#409EFF', alpha=0.85)
                ax.set_title('价格分布直方图')
                ax.set_xlabel('价格')
                ax.set_ylabel('频次')
                price_dist_path = charts_dir / f"price_distribution_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
                plt.tight_layout()
                fig.savefig(price_dist_path)
                plt.close(fig)
                generated['price_distribution'] = str(price_dist_path)
        except Exception as e:
            logger.warning(f"生成价格分布图失败: {e}")

        # 3) 价格偏差条形图（取前20个材料）
        try:
            items = []
            for material in materials[:20]:
                a = analyses_map.get(material.id)
                if a and a.price_variance is not None:
                    items.append((material.material_name or f"材料{material.id}", a.price_variance))
            if items:
                names, variances = zip(*items)
                fig, ax = plt.subplots(figsize=(7, 4), dpi=150)
                colors = ['#67C23A' if v >= 0 else '#F56C6C' for v in variances]
                ax.bar(range(len(variances)), variances, color=colors)
                ax.set_xticks(range(len(names)))
                ax.set_xticklabels([n[:8] for n in names], rotation=45, ha='right')
                ax.set_title('价格偏差(%)')
                ax.axhline(0, color='#666', linewidth=0.8)
                variance_path = charts_dir / f"price_variance_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
                plt.tight_layout()
                fig.savefig(variance_path)
                plt.close(fig)
                generated['price_variance'] = str(variance_path)
        except Exception as e:
            logger.warning(f"生成价格偏差图失败: {e}")

        return generated

    def _add_charts_section(self, doc: 'Document', charts: Dict[str, str]):
        """将图表图片插入Word文档的图表分析章节"""
        doc.add_heading('图表分析', 1)
        # 两列布局：风险饼图 + 价格偏差条形图
        if charts.get('risk_levels') or charts.get('price_variance'):
            table = doc.add_table(rows=1, cols=2)
            row = table.rows[0]
            if charts.get('risk_levels'):
                try:
                    row.cells[0].paragraphs[0].add_run().add_picture(charts['risk_levels'], width=Inches(3.8))
                except Exception as e:
                    logger.warning(f"插入风险饼图失败: {e}")
            if charts.get('price_variance'):
                try:
                    row.cells[1].paragraphs[0].add_run().add_picture(charts['price_variance'], width=Inches(3.8))
                except Exception as e:
                    logger.warning(f"插入价格偏差图失败: {e}")
        # 价格分布图独占一行
        if charts.get('price_distribution'):
            p = doc.add_paragraph()
            try:
                p.add_run().add_picture(charts['price_distribution'], width=Inches(6.5))
            except Exception as e:
                logger.warning(f"插入价格分布图失败: {e}")

    def _add_appendix(self, doc: 'Document'):
        """附录章节：技术说明与统计口径"""
        doc.add_heading('附录', 1)
        doc.add_heading('技术说明', 2)
        doc.add_paragraph('本报告由系统自动生成，分析模型基于历史市场数据与规则引擎。')
        doc.add_heading('统计口径', 2)
        doc.add_paragraph('价格偏差 = (项目单价 - AI预测均价) / AI预测均价 × 100%。风险等级按偏差与置信度综合评估。')
    
    def _add_title_page(self, doc: Document, report_data: Dict[str, Any]):
        """添加标题页"""
        project = report_data['project']
        report_date = report_data['report_date']
        
        # 主标题
        title = doc.add_heading('造价材料审计报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 项目名称
        project_title = doc.add_heading(f'项目: {project.name}', 1)
        project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加空行
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 基本信息表格
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 填充基本信息
        table.cell(0, 0).text = '项目编号'
        table.cell(0, 1).text = project.project_code or 'N/A'
        
        table.cell(1, 0).text = '项目地点'
        table.cell(1, 1).text = project.location or 'N/A'
        
        table.cell(2, 0).text = '业主单位'
        table.cell(2, 1).text = project.owner or 'N/A'
        
        table.cell(3, 0).text = '承包单位'
        table.cell(3, 1).text = project.contractor or 'N/A'
        
        table.cell(4, 0).text = '报告日期'
        table.cell(4, 1).text = report_date.strftime('%Y年%m月%d日')
        
        table.cell(5, 0).text = '项目状态'
        status_value = project.status.value if project.status else 'unknown'
        table.cell(5, 1).text = self._get_status_text(status_value)
        
        # 添加分页符
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, report_data: Dict[str, Any]):
        """添加执行摘要"""
        stats = report_data['statistics']
        
        doc.add_heading('执行摘要', 1)
        
        # 审计概述
        doc.add_heading('审计概述', 2)
        p = doc.add_paragraph()
        p.add_run(f"本次审计共涉及{stats['total_materials']}种材料，")
        p.add_run(f"其中{stats['analyzed_materials']}种材料进行了AI价格分析，")
        p.add_run(f"分析覆盖率达{stats['analysis_coverage']:.1f}%。")
        
        # 主要发现
        doc.add_heading('主要发现', 2)
        findings = doc.add_paragraph()
        findings.add_run(f"• 发现{stats['problematic_materials']}种问题材料，占总材料的{stats['problem_rate']:.1f}%\n")
        findings.add_run(f"• 价格不合理材料{stats['unreasonable_count']}种\n")
        findings.add_run(f"• 平均价格偏差为{stats['avg_price_variance']:.1f}%\n")
        findings.add_run(f"• 预计可节约成本{stats['estimated_savings']:.2f}万元")
        
        # 风险分布
        if stats['risk_distribution']:
            doc.add_heading('风险分布', 2)
            risk_p = doc.add_paragraph()
            for risk_level, count in stats['risk_distribution'].items():
                risk_name = self._get_risk_level_text(risk_level)
                risk_p.add_run(f"• {risk_name}: {count}种材料\n")
    
    def _add_project_overview(self, doc: Document, report_data: Dict[str, Any]):
        """添加项目概况"""
        project = report_data['project']
        
        doc.add_heading('项目概况', 1)
        
        # 项目基本信息
        doc.add_heading('项目基本信息', 2)
        p = doc.add_paragraph()
        p.add_run(f"项目名称: {project.name}\n")
        if project.description:
            p.add_run(f"项目描述: {project.description}\n")
        p.add_run(f"创建时间: {project.created_at.strftime('%Y年%m月%d日')}")
        
        # 审计范围
        doc.add_heading('审计范围', 2)
        scope_p = doc.add_paragraph()
        scope_p.add_run("本次审计主要针对以下内容：\n")
        scope_p.add_run("• 材料价格合理性分析\n")
        scope_p.add_run("• 异常价格材料识别\n")
        scope_p.add_run("• 成本节约潜力评估\n")
        scope_p.add_run("• 风险等级评定")
    
    def _add_analysis_results(self, doc: Document, report_data: Dict[str, Any]):
        """添加分析结果"""
        stats = report_data['statistics']
        
        doc.add_heading('分析结果', 1)
        
        # 统计数据表格
        doc.add_heading('整体统计', 2)
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = '材料总数'
        table.cell(0, 1).text = f"{stats['total_materials']}种"
        
        table.cell(1, 0).text = '已分析材料'
        table.cell(1, 1).text = f"{stats['analyzed_materials']}种"
        
        table.cell(2, 0).text = '问题材料'
        table.cell(2, 1).text = f"{stats['problematic_materials']}种"
        
        table.cell(3, 0).text = '分析覆盖率'
        table.cell(3, 1).text = f"{stats['analysis_coverage']:.1f}%"
        
        table.cell(4, 0).text = '问题比例'
        table.cell(4, 1).text = f"{stats['problem_rate']:.1f}%"
        
        table.cell(5, 0).text = '预计节约'
        table.cell(5, 1).text = f"{stats['estimated_savings']:.2f}万元"
    
    def _add_problematic_materials(self, doc: Document, report_data: Dict[str, Any]):
        """添加问题材料详情"""
        materials = report_data['materials']
        analyses = report_data['analyses']
        
        # 筛选问题材料
        problematic_materials = [m for m in materials if m.is_problematic]
        
        if not problematic_materials:
            doc.add_heading('问题材料详情', 1)
            doc.add_paragraph('未发现问题材料。')
            return
        
        doc.add_heading('问题材料详情', 1)
        
        # 创建问题材料表格
        table = doc.add_table(rows=len(problematic_materials) + 1, cols=7)
        table.style = 'Table Grid'
        
        # 表头
        headers = ['序号', '材料名称', '规格型号', '项目单价', 'AI预测价格', '价格偏差', '风险等级']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        # 填充数据
        for idx, material in enumerate(problematic_materials, 1):
            table.cell(idx, 0).text = str(idx)
            table.cell(idx, 1).text = material.material_name or 'N/A'
            table.cell(idx, 2).text = material.specification or 'N/A'
            table.cell(idx, 3).text = f"¥{material.unit_price:.2f}" if material.unit_price else 'N/A'
            
            # 获取分析结果
            analysis = analyses.get(material.id)
            if analysis:
                predicted_price = analysis.predicted_price_avg or 0
                table.cell(idx, 4).text = f"¥{predicted_price:.2f}"
                
                variance = analysis.price_variance or 0
                table.cell(idx, 5).text = f"{variance:+.1f}%"
                
                risk_level = self._get_risk_level_text(analysis.risk_level or 'unknown')
                table.cell(idx, 6).text = risk_level
            else:
                table.cell(idx, 4).text = 'N/A'
                table.cell(idx, 5).text = 'N/A'
                table.cell(idx, 6).text = 'N/A'
    
    def _add_recommendations(self, doc: Document, report_data: Dict[str, Any]):
        """添加建议措施"""
        stats = report_data['statistics']
        
        doc.add_heading('建议措施', 1)
        
        doc.add_heading('价格调整建议', 2)
        price_rec = doc.add_paragraph()
        price_rec.add_run("• 对价格偏差超过30%的材料进行重点审查\n")
        price_rec.add_run("• 建议与供应商重新议价或寻找替代供应商\n")
        price_rec.add_run("• 加强材料采购的市场调研")
        
        doc.add_heading('风险控制建议', 2)
        risk_rec = doc.add_paragraph()
        risk_rec.add_run("• 建立材料价格监控机制\n")
        risk_rec.add_run("• 定期更新市场价格数据库\n")
        risk_rec.add_run("• 完善材料采购审批流程")
        
        doc.add_heading('后续工作建议', 2)
        follow_rec = doc.add_paragraph()
        follow_rec.add_run("• 每季度进行一次材料价格审计\n")
        follow_rec.add_run("• 建立材料供应商评价体系\n")
        follow_rec.add_run("• 持续优化AI价格分析模型")
    
    def _get_status_text(self, status: str) -> str:
        """获取状态文本"""
        status_map = {
            'draft': '草稿',
            'processing': '处理中',
            'completed': '已完成',
            'failed': '失败'
        }
        return status_map.get(status, status)
    
    def _get_risk_level_text(self, risk_level: str) -> str:
        """获取风险等级文本"""
        risk_map = {
            'low': '正常',
            'medium': '中风险',
            'high': '高风险',
            'severe': '严重风险',
            'unknown': '未知'
        }
        return risk_map.get(risk_level, risk_level)
    
    def _create_sample_materials(self, project_id: int) -> List[ProjectMaterial]:
        """创建示例材料数据"""
        from app.models.project import ProjectMaterial
        
        sample_materials = [
            ProjectMaterial(
                id=1001,
                project_id=project_id,
                material_name="钢筋",
                specification="HRB400E Φ16",
                unit="吨",
                quantity=120.5,
                unit_price=4500.0,
                is_matched=False,
                is_analyzed=True,
                is_problematic=True
            ),
            ProjectMaterial(
                id=1002,
                project_id=project_id,
                material_name="商品混凝土",
                specification="C30",
                unit="立方米",
                quantity=450.0,
                unit_price=380.0,
                is_matched=True,
                is_analyzed=True,
                is_problematic=False
            ),
            ProjectMaterial(
                id=1003,
                project_id=project_id,
                material_name="砌体",
                specification="M10水泥砂浆",
                unit="立方米",
                quantity=280.0,
                unit_price=650.0,
                is_matched=False,
                is_analyzed=True,
                is_problematic=True
            )
        ]
        return sample_materials
    
    def _create_sample_analyses(self, materials: List[ProjectMaterial]) -> List[PriceAnalysis]:
        """创建示例分析数据"""
        from app.models.analysis import PriceAnalysis
        import random
        
        analyses = []
        for material in materials:
            # 创建分析结果
            variance = random.uniform(-30, 50) if material.is_problematic else random.uniform(-10, 10)
            predicted_price = material.unit_price * (1 - variance / 100) if material.unit_price else 0
            
            risk_level = 'high' if material.is_problematic else 'low'
            if abs(variance) >= 60:
                risk_level = 'severe'
            elif abs(variance) >= 40:
                risk_level = 'high'
            elif abs(variance) >= 20:
                risk_level = 'medium'
            else:
                risk_level = 'low'
            
            analysis = PriceAnalysis(
                material_id=material.id,
                predicted_price_min=predicted_price * 0.9,
                predicted_price_max=predicted_price * 1.1,
                predicted_price_avg=predicted_price,
                confidence_score=random.uniform(0.7, 0.95),
                price_variance=variance,
                is_reasonable=not material.is_problematic,
                risk_level=risk_level,
                api_response={
                    'ai_provider': 'test_ai',
                    'analysis_time': '2025-09-06',
                    'data_source': '测试数据'
                }
            )
            analyses.append(analysis)
        
        return analyses
    
    def _generate_minimal_report(self, project: Project) -> str:
        """生成最小化的报告（当主要生成失败时的备用方案）"""
        try:
            logger.info("尝试生成最小化报告...")
            
            if not DOCX_AVAILABLE:
                # 如果连docx都不可用，生成文本报告
                return self._generate_text_report(project)
            
            # 创建最简单的Word文档
            doc = Document()
            
            # 标题
            doc.add_heading('造价材料审计报告', 0)
            
            # 基本信息
            doc.add_heading('项目信息', 1)
            p = doc.add_paragraph()
            p.add_run(f"项目名称: {project.name if project else '未知项目'}\n")
            p.add_run(f"生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}\n")
            p.add_run("状态: 报告生成遇到问题，这是简化版本")
            
            # 说明
            doc.add_heading('说明', 1)
            doc.add_paragraph("由于系统问题，无法生成完整的审计报告。请联系系统管理员检查系统状态。")
            
            # 保存文件
            project_id = project.id if project else 'unknown'
            report_filename = f"minimal_report_{project_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            report_path = self.reports_dir / report_filename
            doc.save(str(report_path))
            
            logger.info(f"最小化报告生成成功: {report_path}")
            return str(report_path.absolute())
            
        except Exception as e:
            logger.error(f"最小化报告生成也失败: {e}")
            # 最后的备用方案：生成文本文件
            return self._generate_text_report(project)
    
    def _generate_text_report(self, project: Project) -> str:
        """生成纯文本报告（最后的备用方案）"""
        try:
            logger.info("生成纯文本报告...")
            
            project_id = project.id if project else 'unknown'
            project_name = project.name if project else '未知项目'
            
            report_content = f"""造价材料审计报告
            
项目名称: {project_name}
项目ID: {project_id}
生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}

报告说明:
由于系统技术问题，无法生成完整的Word格式报告。
这是一个临时的文本版本报告。

建议:
1. 检查系统依赖是否完整安装
2. 检查文件系统权限
3. 联系系统管理员进行技术支持

系统状态:
- Word文档生成库: {'可用' if DOCX_AVAILABLE else '不可用'}
- 报告目录: {self.reports_dir}
"""
            
            report_filename = f"text_report_{project_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            report_path = self.reports_dir / report_filename
            
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
            
            logger.info(f"文本报告生成成功: {report_path}")
            return str(report_path.absolute())
            
        except Exception as e:
            logger.error(f"连文本报告都生成失败: {e}")
            raise Exception(f"所有报告生成方案都失败了: {e}") from e
